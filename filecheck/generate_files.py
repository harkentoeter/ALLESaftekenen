# generate_files.py

from openpyxl import Workbook
from docx import Document

# Sample cybersecurity data
data = [
    {"name": "Alice Hacking", "email": "alice@cybersec.com", "phone": "+31 6 12345678", "address": "Hackerstraat 1, 1337 AB Amsterdam"},
    {"name": "Bob Breach", "email": "bob.breach@infosec.org", "phone": "+31 6 87654321", "address": "Datadiefweg 99, 1000 ZZ Utrecht"},
    {"name": "Eve Phisher", "email": "eve.phisher@malmail.net", "phone": "+31 6 99999999", "address": "Phishinglaan 123, 2000 DD Den Haag"}
]

# Create Word document
doc = Document()
doc.add_heading("Cyber Security Report", 0)

for person in data:
    doc.add_paragraph(
        f"Naam: {person['name']}\n"
        f"Email: {person['email']}\n"
        f"Telefoon: {person['phone']}\n"
        f"Adres: {person['address']}\n"
        f"{'-'*30}"
    )

doc.save("cyber_report.docx")

# Create Excel document
wb = Workbook()
ws = wb.active
ws.title = "CyberData"

# Headers
ws.append(["Naam", "Email", "Telefoon", "Adres"])

# Data
for person in data:
    ws.append([person["name"], person["email"], person["phone"], person["address"]])

wb.save("cyber_data.xlsx")

print("Bestanden succesvol gegenereerd: cyber_report.docx & cyber_data.xlsx")
